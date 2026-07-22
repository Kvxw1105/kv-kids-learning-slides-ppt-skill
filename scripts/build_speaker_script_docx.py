#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SLIDE_RE = re.compile(r"^##\s+(S\d+)\s*[·.-]\s*(.+?)\s*$")


@dataclass(frozen=True)
class Theme:
    accent: str
    accent_dark: str
    secondary: str
    cream: str
    pale: str
    ink: str
    muted: str
    border: str


THEMES = {
    "paper-theater": Theme("E98DA8", "A64B6A", "91A77A", "FFF7E8", "FDE7EC", "3E3034", "7C6B70", "EBC7D2"),
    "watercolor-glow": Theme("8B78C6", "51447A", "6AA6C8", "FFF9EF", "EEEAF9", "343044", "716B82", "D9D1EE"),
    "clean-editorial": Theme("E36F63", "8F3932", "234E70", "FAFAF7", "EEF3F7", "25333D", "64737D", "CFD9E0"),
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def hex_rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str, size: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepLines")
    p_pr.append(node)


def set_run_font(run, east_asia: str, latin: str | None = None, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = latin or east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = hex_rgb(color)


def add_text(paragraph, text: str, *, size: float, bold: bool = False, color: str, east_asia: str = "Noto Sans CJK SC", latin: str = "Liberation Sans"):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia, latin, size, bold, color)
    return run


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def parse_markdown(path: Path) -> dict:
    lines = path.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "Speaker Script")
    metadata: list[tuple[str, str]] = []
    slides: list[dict] = []
    usage: list[str] = []
    continuous: dict[str, list[str]] = {"english": [], "chinese": []}
    current = None
    subsection = None
    mode = None
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("> **") and "：**" in line:
            key, value = line[4:].split("：**", 1)
            metadata.append((key.strip(), value.strip()))
            continue
        if line.startswith("## 使用建议") or line.startswith("## How to Use"):
            mode = "usage"
            current = None
            subsection = None
            continue
        if line.startswith("## Continuous English") or line.startswith("## 英文连贯稿"):
            mode = "continuous_en"
            current = None
            subsection = None
            continue
        if line.startswith("## 中文参考稿") or line.startswith("## Continuous Chinese"):
            mode = "continuous_zh"
            current = None
            subsection = None
            continue
        match = SLIDE_RE.match(line)
        if match:
            current = {"id": match.group(1), "title": match.group(2), "sections": {}}
            slides.append(current)
            subsection = None
            mode = "slide"
            continue
        if line.startswith("### ") and current:
            subsection = line[4:].strip()
            current["sections"][subsection] = []
            continue
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue
        if mode == "usage" and stripped.startswith("-"):
            usage.append(stripped.lstrip("- "))
        elif mode == "continuous_en":
            continuous["english"].append(stripped)
        elif mode == "continuous_zh":
            continuous["chinese"].append(stripped)
        elif current and subsection:
            current["sections"][subsection].append(stripped)
    return {"title": title, "metadata": metadata, "usage": usage, "slides": slides, "continuous": continuous}


def get_section(slide: dict, prefixes: tuple[str, ...]) -> list[str]:
    for key, values in slide.get("sections", {}).items():
        if any(key.lower().startswith(prefix.lower()) for prefix in prefixes):
            return values
    return []


def clean_paragraph_lines(values: list[str]) -> list[str]:
    return [value.lstrip("- ").lstrip("> ").strip() for value in values if value.strip()]


def configure_document(document: Document, theme: Theme, source_path: Path) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = hex_rgb(theme.ink)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color in (("Title", 28, theme.accent_dark), ("Heading 1", 20, theme.accent_dark), ("Heading 2", 15, theme.accent_dark), ("Heading 3", 11.5, theme.secondary)):
        style = document.styles[style_name]
        style.font.name = "Liberation Sans"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = hex_rgb(color)

    document.core_properties.title = "Bilingual Child Speech Script"
    document.core_properties.subject = "Editable bilingual child speech script generated from Markdown"
    document.core_properties.author = "kv-kids-learning-slides"
    document.core_properties.comments = f"source_markdown_sha256={sha256(source_path)}"


def add_header_footer(document: Document, title: str, theme: Theme) -> None:
    for section in document.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, title, size=8.5, bold=True, color=theme.muted)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(fp, "CHILD SPEECH SCRIPT  |  ", size=8, bold=True, color=theme.muted)
        add_field(fp, "PAGE")


def add_cover(document: Document, data: dict, theme: Theme) -> None:
    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    cell = band.cell(0, 0)
    cell.width = Inches(6.9)
    set_cell_shading(cell, theme.accent)
    set_cell_margins(cell, top=40, bottom=40)
    cell.paragraphs[0].add_run(" ")

    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, data["title"], size=27, bold=True, color=theme.accent_dark, east_asia="Noto Serif CJK SC", latin="Liberation Serif")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "Bilingual Child Speech Script", size=13, bold=True, color=theme.secondary)
    add_text(subtitle, "  /  儿童双语演讲推荐稿", size=13, bold=True, color=theme.secondary)

    document.add_paragraph()
    meta = document.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    chosen = data["metadata"][:8]
    if not chosen:
        chosen = [("Speaker / 演讲者", "________"), ("Duration / 建议时长", "________")]
    for index, (key, value) in enumerate(chosen):
        row = meta.add_row()
        row.cells[0].width = Inches(2.35)
        row.cells[1].width = Inches(4.15)
        for cell in row.cells:
            set_cell_margins(cell, top=125, bottom=125, start=170, end=170)
            set_cell_border(cell, theme.border, 6)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], theme.pale if index % 2 == 0 else theme.cream)
        set_cell_shading(row.cells[1], "FFFFFF")
        p0 = row.cells[0].paragraphs[0]
        add_text(p0, key, size=9.5, bold=True, color=theme.accent_dark)
        p1 = row.cells[1].paragraphs[0]
        add_text(p1, value, size=10.5, bold=False, color=theme.ink)

    document.add_paragraph()
    note = document.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.autofit = False
    note_cell = note.cell(0, 0)
    set_cell_shading(note_cell, theme.cream)
    set_cell_border(note_cell, theme.border, 7)
    set_cell_margins(note_cell, top=180, bottom=180, start=220, end=220)
    p = note_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Editable source: Markdown  ·  Fully editable Word packaging", size=10, bold=True, color=theme.secondary)
    p2 = note_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p2, "建议先修改 .md 内容，再重新生成 .docx，避免两个版本不一致。", size=9.5, color=theme.muted)
    document.add_page_break()


def add_usage(document: Document, data: dict, theme: Theme) -> None:
    heading = document.add_heading("使用建议 / How to Use", level=1)
    keep_with_next(heading)
    tips = data["usage"] or [
        "先用英文稿练熟故事顺序，再借助中文理解意思。",
        "动作提示用于排练，不需要逐字念出。",
        "家长或老师可以直接修改本 Word，也可以先修改 Markdown 再重新生成。",
    ]
    box = document.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    set_cell_shading(cell, theme.pale)
    set_cell_border(cell, theme.border, 6)
    set_cell_margins(cell, top=160, bottom=150, start=220, end=220)
    for index, tip in enumerate(tips):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.style = document.styles["Normal"]
        add_text(p, f"{index + 1}. ", size=10, bold=True, color=theme.accent_dark)
        add_text(p, tip, size=10, color=theme.ink)
    document.add_paragraph()


def add_slide_section(document: Document, slide: dict, theme: Theme, index: int) -> None:
    header = document.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(0.8)
    header.columns[1].width = Inches(5.9)
    left, right = header.rows[0].cells
    set_cell_shading(left, theme.accent)
    set_cell_shading(right, theme.cream)
    for cell in (left, right):
        set_cell_border(cell, theme.border, 7)
        set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(lp, slide["id"], size=10.5, bold=True, color="FFFFFF")
    rp = right.paragraphs[0]
    add_text(rp, slide["title"], size=13, bold=True, color=theme.accent_dark)

    bilingual = document.add_table(rows=2, cols=2)
    bilingual.alignment = WD_TABLE_ALIGNMENT.CENTER
    bilingual.autofit = False
    bilingual.columns[0].width = Inches(3.35)
    bilingual.columns[1].width = Inches(3.35)
    headers = bilingual.rows[0].cells
    set_repeat_table_header(bilingual.rows[0])
    for cell, label, fill, color in (
        (headers[0], "ENGLISH", theme.accent_dark, "FFFFFF"),
        (headers[1], "中文参考", theme.secondary, "FFFFFF"),
    ):
        set_cell_shading(cell, fill)
        set_cell_border(cell, theme.border, 6)
        set_cell_margins(cell, top=70, bottom=70, start=150, end=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, label, size=9.5, bold=True, color=color)
    body = bilingual.rows[1].cells
    en_lines = clean_paragraph_lines(get_section(slide, ("English",)))
    zh_lines = clean_paragraph_lines(get_section(slide, ("中文", "Chinese")))
    for cell, values, fill, lang in (
        (body[0], en_lines, "FFFFFF", "en"),
        (body[1], zh_lines, theme.cream, "zh"),
    ):
        set_cell_shading(cell, fill)
        set_cell_border(cell, theme.border, 6)
        set_cell_margins(cell, top=170, bottom=170, start=190, end=190)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell.text = ""
        values = values or ["[Editable text]"]
        for idx, value in enumerate(values):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            p.paragraph_format.line_spacing = 1.22
            p.paragraph_format.space_after = Pt(3)
            keep_together(p)
            add_text(p, value, size=10.5 if lang == "en" else 10.2, color=theme.ink, east_asia="Noto Sans CJK SC", latin="Liberation Sans")

    actions = clean_paragraph_lines(get_section(slide, ("Stage actions", "舞台动作", "Actions")))
    rhythm = clean_paragraph_lines(get_section(slide, ("Rhythm", "节奏", "Pronunciation")))
    notes = clean_paragraph_lines(get_section(slide, ("Editable notes", "可修改备注", "Notes")))

    cues = document.add_table(rows=1, cols=2)
    cues.alignment = WD_TABLE_ALIGNMENT.CENTER
    cues.autofit = False
    cues.columns[0].width = Inches(3.35)
    cues.columns[1].width = Inches(3.35)
    for cell, label, values, fill in (
        (cues.cell(0, 0), "STAGE ACTIONS / 舞台动作", actions or ["No special action / 无特别动作"], theme.pale),
        (cues.cell(0, 1), "RHYTHM & PRONUNCIATION / 节奏与发音", rhythm or ["Speak slowly and pause naturally / 慢一点，自然停顿"], theme.cream),
    ):
        set_cell_shading(cell, fill)
        set_cell_border(cell, theme.border, 6)
        set_cell_margins(cell, top=120, bottom=120, start=170, end=170)
        p = cell.paragraphs[0]
        add_text(p, label, size=8.8, bold=True, color=theme.accent_dark)
        for value in values:
            bp = cell.add_paragraph(style=None)
            bp.paragraph_format.left_indent = Cm(0.15)
            add_text(bp, "• ", size=9.2, bold=True, color=theme.secondary)
            add_text(bp, value, size=9.2, color=theme.ink)

    note_table = document.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    note_cell = note_table.cell(0, 0)
    set_cell_shading(note_cell, "FFFFFF")
    set_cell_border(note_cell, theme.border, 6)
    set_cell_margins(note_cell, top=90, bottom=90, start=170, end=170)
    p = note_cell.paragraphs[0]
    add_text(p, "EDITABLE NOTES / 可修改备注", size=8.6, bold=True, color=theme.muted)
    note_text = " ".join(notes) if notes else "____________________________________________________________"
    np = note_cell.add_paragraph()
    add_text(np, note_text, size=9, color=theme.muted)
    document.add_paragraph().paragraph_format.space_after = Pt(2)




def add_final_checklist(document: Document, theme: Theme) -> None:
    heading = document.add_heading("上台前最后检查 / Final Stage Check", level=2)
    keep_with_next(heading)
    table = document.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    items = [
        ("1", "Stand tall and smile / 站稳，微笑"),
        ("2", "Look at the audience / 看一看观众"),
        ("3", "Speak slowly and clearly / 慢一点，说清楚"),
        ("4", "Pause between story beats / 每一段之间停一下"),
        ("5", "Use actions gently / 动作自然，不要太快"),
        ("6", "Finish with a bow or wave / 最后鞠躬或挥手"),
    ]
    for index, (number, text) in enumerate(items):
        cell = table.cell(index // 2, index % 2)
        set_cell_shading(cell, theme.pale if index % 2 == 0 else theme.cream)
        set_cell_border(cell, theme.border, 6)
        set_cell_margins(cell, top=125, bottom=125, start=160, end=160)
        p = cell.paragraphs[0]
        add_text(p, f"{number}. ", size=9.5, bold=True, color=theme.accent_dark)
        add_text(p, text, size=9.5, color=theme.ink)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    add_text(note, "My personal reminder / 我的个性提醒：", size=9.2, bold=True, color=theme.muted)
    add_text(note, "__________________________________________________", size=9.2, color=theme.muted)

def add_continuous_copy(document: Document, data: dict, theme: Theme) -> None:
    document.add_page_break()
    heading = document.add_heading("连贯排练稿 / Continuous Rehearsal Copy", level=1)
    keep_with_next(heading)
    en = " ".join(data["continuous"]["english"]).strip()
    zh = " ".join(data["continuous"]["chinese"]).strip()
    if not en:
        en = "\n\n".join(" ".join(clean_paragraph_lines(get_section(slide, ("English",)))) for slide in data["slides"])
    if not zh:
        zh = "\n\n".join(" ".join(clean_paragraph_lines(get_section(slide, ("中文", "Chinese")))) for slide in data["slides"])
    for label, text, fill, color in (
        ("ENGLISH FULL SCRIPT", en, "FFFFFF", theme.accent_dark),
        ("中文参考全文", zh, theme.cream, theme.secondary),
    ):
        table = document.add_table(rows=2, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header, body = table.rows[0].cells[0], table.rows[1].cells[0]
        set_cell_shading(header, color)
        set_cell_shading(body, fill)
        for cell in (header, body):
            set_cell_border(cell, theme.border, 6)
            set_cell_margins(cell, top=120, bottom=120, start=190, end=190)
        hp = header.paragraphs[0]
        add_text(hp, label, size=10, bold=True, color="FFFFFF")
        body.text = ""
        for index, paragraph_text in enumerate([part.strip() for part in text.split("\n\n") if part.strip()]):
            p = body.paragraphs[0] if index == 0 else body.add_paragraph()
            p.paragraph_format.line_spacing = 1.24
            p.paragraph_format.space_after = Pt(7)
            add_text(p, paragraph_text, size=10.2, color=theme.ink)
        document.add_paragraph()


def build_docx_from_markdown(markdown_path: Path, output_path: Path, theme_name: str = "paper-theater") -> None:
    theme = THEMES.get(theme_name, THEMES["paper-theater"])
    data = parse_markdown(markdown_path)
    document = Document()
    configure_document(document, theme, markdown_path)
    add_header_footer(document, data["title"], theme)
    add_cover(document, data, theme)
    add_usage(document, data, theme)
    for index, slide in enumerate(data["slides"], start=1):
        if index > 1 and index % 2 == 1:
            document.add_page_break()
        add_slide_section(document, slide, theme, index)
    add_final_checklist(document, theme)
    add_continuous_copy(document, data, theme)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Compile an editable bilingual speaker-script DOCX from canonical Markdown.")
    parser.add_argument("markdown")
    parser.add_argument("--out", required=True)
    parser.add_argument("--theme", choices=sorted(THEMES), default="paper-theater")
    args = parser.parse_args()
    markdown = Path(args.markdown)
    output = Path(args.out)
    if not markdown.exists():
        raise SystemExit(f"Markdown does not exist: {markdown}")
    build_docx_from_markdown(markdown, output, args.theme)
    print(f"Wrote {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
