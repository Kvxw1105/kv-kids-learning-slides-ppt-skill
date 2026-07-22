#!/usr/bin/env python3
import argparse
import hashlib
import json
import re
import zipfile
from pathlib import Path

from docx import Document


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--markdown")
    args = parser.parse_args()
    path = Path(args.docx)
    errors = []
    warnings = []
    if not path.exists():
        result = {"status": "BLOCKED", "errors": [f"missing file: {path}"], "warnings": []}
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 1
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names:
                errors.append("missing word/document.xml")
            xml_text = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in names
                if name.endswith(".xml") or name.endswith(".rels")
            )
            if "/mnt/data/" in xml_text or "\\Users\\" in xml_text or "/tmp/" in xml_text:
                errors.append("local path leaked into DOCX package")
    except zipfile.BadZipFile:
        errors.append("invalid DOCX ZIP package")
        xml_text = ""
    try:
        document = Document(path)
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        all_text = "\n".join(paragraphs + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
        if not document.core_properties.title:
            warnings.append("core title is empty")
        comments = document.core_properties.comments or ""
        if "source_markdown_sha256=" not in comments:
            errors.append("missing source Markdown SHA-256 binding")
        elif args.markdown:
            markdown_path = Path(args.markdown)
            if not markdown_path.exists():
                errors.append(f"source Markdown missing: {markdown_path}")
            else:
                h = hashlib.sha256(markdown_path.read_bytes()).hexdigest()
                if f"source_markdown_sha256={h}" not in comments:
                    errors.append("DOCX source hash does not match current Markdown")
        if len(document.tables) < 4:
            errors.append(f"expected structured editable tables, found {len(document.tables)}")
        if "ENGLISH" not in all_text or "中文" not in all_text:
            errors.append("bilingual section labels missing")
        slide_sections = len(re.findall(r"\bS\d{2}\b", all_text))
        if slide_sections < 1:
            errors.append("no Sxx slide sections found")
        if "Continuous Rehearsal Copy" not in all_text and "连贯排练稿" not in all_text:
            warnings.append("continuous rehearsal section not detected")
    except Exception as exc:
        errors.append(f"python-docx read failed: {exc}")
        paragraphs = []
        slide_sections = 0
    status = "BLOCKED" if errors else ("PASS_WITH_WARNINGS" if warnings else "PASS")
    result = {
        "status": status,
        "errors": errors,
        "warnings": warnings,
        "size_bytes": path.stat().st_size,
        "paragraphs": len(paragraphs),
        "tables": len(document.tables) if 'document' in locals() else 0,
        "slide_sections_detected": slide_sections,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 1 if errors else 0


if __name__ == "__main__":
    raise SystemExit(main())
